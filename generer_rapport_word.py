from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Titre principal
titre = doc.add_heading('Rapport de stage en responsabilité', level=0)
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('RÉPUBLIQUE DU NIGER  UNIVERSITÉ ABDOU MOUMOUNI  ÉCOLE NORMALE SUPÉRIEURE')
doc.add_paragraph('MPPES2 Option : Philosophie')
doc.add_paragraph()
doc.add_paragraph('Rapport de stage en responsabilité pour l’obtention de Master Professionnel au Professorat de l’Enseignement Secondaire (MPPES)')
doc.add_paragraph('Effectué du 05 Mars au 30 Avril 2024 au CES/Gaweye1.')
doc.add_paragraph('Rédigé par : Daouda Moussa Sani N° de carte : 82279')
doc.add_paragraph('Encadré par : Inspecteur Adamou Moussa Dr Boubakar Maizoumbou')
doc.add_paragraph('Tutrice : Ibrahim Hassane Nana Zoulaha')
doc.add_paragraph('Année académique : 2023-2024')
doc.add_page_break()

# Table des matières (manuelle)
doc.add_heading('Table des matières', level=1)
tdm = [
    "Remerciements", "Sigles et abréviations", "Introduction générale",
    "1. Présentation de l’établissement d’accueil",
    "1-1. Historique de l’établissement",
    "1-2. Situation du personnel administratif et autre",
    "1-3. Situation du personnel enseignant",
    "1-3-1. Situation par discipline et sexe",
    "1-3-2. Situation par statut et sexe",
    "1-3-3. Situation par grade et sexe",
    "1-4. Situation des élèves",
    "1-4-1. Premier cycle par niveau et sexe",
    "1-4-2. Second cycle par niveau et sexe",
    "1-5. Situation des infrastructures disponibles",
    "II. Activités réalisées",
    "III. Constats et réflexions personnels",
    "Conclusion générale et recommandations",
    "Références bibliographiques",
    "Annexes"
]
for item in tdm:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# Remerciements
doc.add_heading('Remerciements', level=1)
remerciements = """Tout d’abord je remercie le bon Dieu, qui par sa grâce infinie m’a accordé le temps et la santé pour réaliser ce stage de fin de formation pour le MPPES2.

Mes sincères remerciements vont également à l’ensemble des enseignants-chercheurs de l’ENS (Ecole Normale Supérieure) et particulièrement ceux de philosophie qui m’ont assuré avec dévouement et abnégation une formation combien précieuse pour ma carrière d’enseignant, qu’ils trouvent ici toute ma reconnaissance.

Ensuite, je remercie très sincèrement l’inspecteur Adamou Moussa et le docteur Boubacar Maïzoumbou pour m’avoir encadré sur le terrain durant toute la période du stage, en mettant leurs expériences et expertises d’enseignement à ma disposition afin que j’améliore mes pratiques enseignantes.

Mes remerciements vont aussi à l’administration et à l’ensemble du personnel du CES/Gaweye1 notamment le proviseur, le censeur et ma tutrice de stage Ibrahim Hassane Nana Zoulaha pour leur franche collaboration qui m’a permis de réussir mon stage.

Enfin, je remercie ma femme et mes enfants qui m’ont soutenu par leurs prières et supporté mon absence durant tout ce temps de la formation."""
doc.add_paragraph(remerciements)
doc.add_page_break()

# Sigles et abréviations
doc.add_heading('Sigles et abréviations', level=1)
sigles = [
    "CEG : collège d’enseignement général",
    "CES : Complexe d’Enseignement Secondaire",
    "CM : Cours Magistraux",
    "DDEN : Direction Départementale de l’Education Nationale",
    "DREN : Direction Régionale de l’Education Nationale",
    "ENS : Ecole Normale Supérieure",
    "IESG : Inspection d’Enseignement Secondaire Général",
    "IPR : Inspection Pédagogique Régionale",
    "LEG : Lycée d’Enseignement Général",
    "LPPES : Licence Professionnel au Professorat de l’Enseignement Secondaire",
    "MEN: Ministère de l’Education Nationale",
    "MPPES : Master Professionnel au Professorat de l’Enseignement Secondaire",
    "PES : Professeur d’Enseignement Secondaire",
    "TD : Travaux Dirigés",
    "UAM : Université Abdou Moumouni",
    "UP : Unité Pédagogique"
]
for s in sigles:
    doc.add_paragraph(s, style='List Bullet')
doc.add_page_break()

# Introduction générale
doc.add_heading('Introduction générale', level=1)
intro = """C’est dans le but d’améliorer la qualité de l’enseignement/apprentissage au second cycle du secondaire, en collaboration avec le Ministère de l’Education Nationale, de l’Alphabétisation, de l’Enseignement Professionnel et de la Promotion des Langues Nationales (MEN/A/EP/PLN) et le Ministère de la Fonction Publique et de la Réforme Administrative (MFP/RA), que l’Ecole Normale Supérieure (ENS) a institué la formation de Master Professionnel au Professorat de l’Enseignement Secondaire (MPPES2). En effet, deux phases caractérisent cette formation : la première étant essentiellement consacrée aux Cours Magistraux (CM) et Travaux Dirigés (TD) ; la seconde au stage. Après le stage d’observation, le stage en responsabilité a pour objectif de mettre le stagiaire au contact de la réalité et de lui permettre d’appliquer les connaissances acquises tout au long de la phase théorique. C’est en ce sens que le stage a été officiellement lancé à travers la correspondance, N°000022/DREN/A/EP/PLN/2023, du 19/01/2024, du Directeur régional de Niamey, adressée aux inspecteurs de l’enseignement secondaire général/FA.

Ainsi, du 05 Mars au 30 Avril 2024, période correspondant au stage en responsabilité, j’ai été mis à la disposition du CES/Gaweye1, avec comme tutrice Ibrahim Hassane Nana Zoulaha, titulaire d’un MPPES2 en Philosophie, obtenu en 2023.

Comment s’est déroulé ce stage en responsabilité ? Quels enseignements ai-je retenus de ce stage ? Tout au long de mon travail, je m’attacherai à présenter d’abord mon établissement d’accueil, à exposer ensuite les activités que j’ai eues à réaliser, et à exprimir enfin mes constats et réflexions personnels avant de tirer une conclusion générale et de faire des recommandations."""
doc.add_paragraph(intro)
doc.add_page_break()

# 1. Présentation de l'établissement
doc.add_heading('1. Présentation de l’établissement d’accueil', level=1)
doc.add_heading('1-1. Historique de l’établissement', level=2)
doc.add_paragraph("Le CEG/Gaweye a été créé le 10 Août 1992, par arrêté N°140/MEN-R/DEST/DEP du 10/08/1992. Il a été transformé en CES, le 20/09/2013. Situé dans le quartier Gaweye, il devient CES/Gaweye1, le 05/09/2017. Il est rattaché à la DDEN et l’IESG Niamey 5.")

# Tableau administratif corrigé
doc.add_heading('1-2. Situation du personnel administratif et autre', level=2)
data_admin = [
    ("Proviseur", 0, 1),
    ("Censeur", 0, 1),
    ("Surveillant", 0, 1),
    ("Bibliothécaires", 0, 2),
    ("Agent de bureau", 2, 7),
    ("Technicien de laboratoire", 0, 0),
    ("Gardien", 0, 2),
    ("Planton", 0, 0),
    ("Manœuvre", 0, 0),
    ("Menuisier", 0, 0),
    ("Infirmier", 0, 0),
    ("Informaticien", 0, 0)
]
table_admin = doc.add_table(rows=len(data_admin)+1, cols=3)
table_admin.style = 'Table Grid'
hdr = table_admin.rows[0].cells
hdr[0].text = 'Fonctions'
hdr[1].text = 'Homme'
hdr[2].text = 'Femme'
for i, (fonc, h, f) in enumerate(data_admin):
    row = table_admin.rows[i+1]
    row.cells[0].text = fonc
    row.cells[1].text = str(h)
    row.cells[2].text = str(f)
doc.add_paragraph("Source : le Proviseur du CES/Gaweye1. À partir de ce tableau, l'on constate l'absence de planton et de manœuvre.")

# Enseignant par discipline
doc.add_heading('1-3. Situation du personnel enseignant', level=2)
doc.add_heading('1-3-1. Situation du personnel enseignant par discipline et sexe', level=3)
table_discipline = doc.add_table(rows=3, cols=12)
table_discipline.style = 'Table Grid'
disciplines = ["Français", "Anglais", "H/G", "SVT", "Maths", "EF", "EPS", "PC", "Philo", "Arabe", "Total"]
h1 = table_discipline.rows[0].cells
h1[0].text = 'Sexe'
for idx, d in enumerate(disciplines):
    h1[idx+1].text = d

# Ligne Homme
row_homme = table_discipline.rows[1].cells
row_homme[0].text = 'Homme'
valeurs_homme = ['5', '4', '4', '2', '6', '0', '3', '5', '1', '2', '32']
for idx, val in enumerate(valeurs_homme, start=1):
    row_homme[idx].text = val

# Ligne Femme
row_femme = table_discipline.rows[2].cells
row_femme[0].text = 'Femme'
valeurs_femme = ['5', '8', '4', '3', '4', '3', '2', '1', '2', '0', '32']
for idx, val in enumerate(valeurs_femme, start=1):
    row_femme[idx].text = val

doc.add_paragraph("De cette situation, on remarque que le nombre des hommes est égal à celui des femmes.")

# Par statut
doc.add_heading('1-3-2. Situation du personnel enseignant par statut et sexe', level=3)
table_statut = doc.add_table(rows=3, cols=6)
table_statut.style = 'Table Grid'
statuts = ["Titulaires", "Contractuels", "ASCN", "Autres", "Total"]
h_stat = table_statut.rows[0].cells
h_stat[0].text = 'Sexe'
for i, s in enumerate(statuts):
    h_stat[i+1].text = s

row_h_stat = table_statut.rows[1].cells
row_h_stat[0].text = 'Homme'
valeurs_h_stat = ['21', '0', '9', '0', '30']
for idx, val in enumerate(valeurs_h_stat, start=1):
    row_h_stat[idx].text = val

row_f_stat = table_statut.rows[2].cells
row_f_stat[0].text = 'Femme'
valeurs_f_stat = ['17', '1', '3', '0', '21']
for idx, val in enumerate(valeurs_f_stat, start=1):
    row_f_stat[idx].text = val

doc.add_paragraph("Sur les 64 enseignants que compte le CES/Gaweye1, les titulaires sont plus nombreux : 38 sur 64 (59%).")

# Par grade
doc.add_heading('1-3-3. Situation du personnel enseignant par grade et sexe', level=3)
table_grade = doc.add_table(rows=3, cols=9)
table_grade.style = 'Table Grid'
grades = ["PES", "CE", "P/CEG", "M/EPS", "M/EFS", "B2", "IIA", "Total"]
h_grade = table_grade.rows[0].cells
h_grade[0].text = 'Sexe'
for i, g in enumerate(grades):
    h_grade[i+1].text = g

row_h_grade = table_grade.rows[1].cells
row_h_grade[0].text = 'Homme'
valeurs_h_grade = ['4', '19', '4', '2', '0', '1', '0', '30']
for idx, val in enumerate(valeurs_h_grade, start=1):
    row_h_grade[idx].text = val

row_f_grade = table_grade.rows[2].cells
row_f_grade[0].text = 'Femme'
valeurs_f_grade = ['1', '22', '6', '2', '2', '0', '1', '34']
for idx, val in enumerate(valeurs_f_grade, start=1):
    row_f_grade[idx].text = val

doc.add_paragraph("La lecture de ce tableau nous renseigne que le nombre des enseignants qui ont le grade de C.E est le plus élevé : 41 sur 64 (64,06%).")

# Élèves 1er cycle
doc.add_heading('1-4. Situation des élèves', level=2)
doc.add_heading('1-4-1. Situation des élèves du premier cycle par niveau et sexe', level=3)
table_premier = doc.add_table(rows=3, cols=5)
table_premier.style = 'Table Grid'
niveaux = ["6ème", "5ème", "4ème", "3ème", "Total"]
h_prem = table_premier.rows[0].cells
h_prem[0].text = 'Sexe'
for i, n in enumerate(niveaux):
    h_prem[i+1].text = n

row_f_prem = table_premier.rows[1].cells
row_f_prem[0].text = 'Filles'
valeurs_f_prem = ['360', '182', '123', '126', '791']
for idx, val in enumerate(valeurs_f_prem, start=1):
    row_f_prem[idx].text = val

row_g_prem = table_premier.rows[2].cells
row_g_prem[0].text = 'Garçons'
valeurs_g_prem = ['302', '115', '82', '64', '563']
for idx, val in enumerate(valeurs_g_prem, start=1):
    row_g_prem[idx].text = val

# Élèves 2nd cycle
doc.add_heading('1-4-2. Situation des élèves du second cycle par niveau et sexe', level=3)
table_second = doc.add_table(rows=3, cols=10)
table_second.style = 'Table Grid'
second_niveaux = ["2nde A", "2nde C", "1ère A", "1ère C", "1ère D", "Tle A", "Tle C", "Tle D", "Total"]
h_sec = table_second.rows[0].cells
h_sec[0].text = 'Sexe'
for i, n in enumerate(second_niveaux):
    h_sec[i+1].text = n

row_f_sec = table_second.rows[1].cells
row_f_sec[0].text = 'Filles'
valeurs_f_sec = ['98', '4', '169', '0', '0', '219', '6', '0', '496']
for idx, val in enumerate(valeurs_f_sec, start=1):
    row_f_sec[idx].text = val

row_g_sec = table_second.rows[2].cells
row_g_sec[0].text = 'Garçons'
valeurs_g_sec = ['72', '6', '53', '1', '0', '42', '4', '4', '182']
for idx, val in enumerate(valeurs_g_sec, start=1):
    row_g_sec[idx].text = val

doc.add_paragraph("En faisant la somme des effectifs des deux cycles, le CES/Gaweye1 compte en cette année académique 2023-2024, 1999 élèves dont 1139 filles et 860 garçons, répartis dans 26 classes. L’analyse de ces effectifs nous montre que les filles sont plus nombreuses au 1er comme au 2nd cycles : 56,97%.")

# Infrastructures
doc.add_heading('1-5. Situation des infrastructures disponibles', level=2)
infras = [
    ("Classes en matériaux définitifs", "26"),
    ("Bibliothèque", "01"),
    ("Laboratoire", "01"),
    ("Magasin", "02"),
    ("Infirmerie", "00"),
    ("Bloc administratif", "01"),
    ("Logement proviseur/Censeur/Surveillant", "00"),
    ("Logement Gardien", "1"),
    ("Salle de professeurs", "01"),
    ("Salle internet", "00"),
    ("Terrain de sport", "01")
]
table_infra = doc.add_table(rows=len(infras), cols=2)
table_infra.style = 'Table Grid'
for i, (nom, nb) in enumerate(infras):
    row = table_infra.rows[i]
    row.cells[0].text = nom
    row.cells[1].text = nb
doc.add_paragraph("Il ressort de ce tableau qu’il n’existe de logement ni pour le proviseur, ni pour le censeur et encore moins pour le surveillant.")
doc.add_page_break()

# II. Activités
doc.add_heading('II. Activités réalisées', level=1)
activites = """J’ai fait ma première séance de présentation de la leçon en TA2, Mardi, le 20 Février 2024, de 08h à 10h. Cette leçon a porté sur le chapitre Histoire, elle a pour titre : la question d’écriture et d’oralité en histoire. Après le prérequis, nous avons motivé les apprenants afin d’arriver à la leçon du jour. Puis s’en est suivi l’annonce des objectifs spécifiques qui étaient pour cette leçon au nombre de deux (2) : à la fin de la leçon, les élèves doivent être capables de : préciser que l’écriture est la source de l’histoire ; citer l’oralité comme une autre source en histoire.

Pour atteindre le premier objectif, nous avons procédé au questionnement. Ainsi, nous avons demandé aux élèves de préciser l’importance de l’écriture. Ils ont répondu que l’écriture sert à documenter, conserver, transmettre, vérifier les faits, les idées de ceux qui ont vécu avant nous, assouplir la mémoire. Nous avons bien apprécié ces réponses et un échange d’approfondissement de celles-ci a été fait avec les élèves. Par la suite, nous avons procédé à la dictée de la trace écrite. Après l’épuisement de ce premier objectif spécifique, nous avons débuté le deuxième en demandant aux élèves si l’écriture est la seule source en histoire. Après un temps de réflexion (5mn), les élèves ont donné des réponses relatives à l’oralité, en donnant l’exemple de griot en Afrique et de la citation de Hampaté Bâ. Après l’approfondissement des réponses données par les élèves, nous avons dicté la trace écrite. Avant de poser deux questions d’évaluation, nous avons effacé le tableau. À la question quelles sont les deux sources en histoire, ils ont répondu que c’est l’écriture et l’oralité. À la question de savoir quelle est leur importance, les apprenants ont précisé que les deux permettent à l’historien d’avoir des informations pour comprendre et interpréter le passé. Après ces réponses, nous avons eu un sentiment de satisfaction car les objectifs de la leçon ont été atteints. Les élèves ont fait preuve d’une discipline et ont manifesté la joie d’avoir compris le cours. Mais nous avons remarqué que ces élèves n’avaient pas l’habitude de demander la parole avant de répondre aux questions, ce qui les pousse à répondre tous au même moment. Ce qui nous a donné de la peine à leur faire comprendre qu’il faut demander la parole avant de répondre à une question posée par l’enseignant. (Voir la fiche de leçon N°1 en annexe)."""
doc.add_paragraph(activites)
doc.add_paragraph("Jeudi, le 29 Février 2024, j’ai présenté une leçon du chapitre Travail, en TA2, de 08h à 10h, qui a comme titre : Travail comme source de liberté et activité propre à l’homme. Cette leçon vise deux objectifs spécifiques : justifier que le travail est libérateur et préciser que le travail est spécifique à l’homme. (Cf. la fiche de leçon N°04 en annexe).")
doc.add_paragraph("Jeudi, le 29 Février 2024, de 12h 30mn à 13h 30mn, en TA2, j’ai présenté une leçon titrée : le machinisme et ses conséquences, du chapitre Travail. L’objectif spécifique de cette leçon est ainsi formulé : À la fin de la leçon, les élèves doivent être capables d’argumenter les avantages du machinisme. (Voir la fiche de leçon N°05 en annexe).")
doc.add_paragraph("Mercredi, le 06 Mars 2024, j’ai eu à présenter une leçon du chapitre Travail dont le titre est : le machinisme et ses conséquences. Déroulée de 11h 30mn à 12h 30mn, en TA2, cette leçon avait comme objectif spécifique : À la fin de la leçon, les élèves doivent être capables de préciser les inconvénients du machinisme. (Cf. la fiche de leçon N°06 en annexe). La présentation de cette leçon a été faite sous la supervision du Docteur Boubacar Maizoumbou, accompagné de ma tutrice. À la fin de cette présentation de la leçon, la phase d’entretien a permis au Dr Boubacar de faire des observations touchant la forme et le fond de ma prestation. Du point de vue forme, le Docteur a relevé que les élèves étaient motivés, la classe a été maîtrisée, les élèves ont bien participé et le tableau est bien divisé. En ce qui concerne le fond, le Docteur a mentionné que les prérequis ont bien été rappelés, le texte-support bien choisi, il a également noté que les élèves ont donné des réponses correctes et que l’évaluation a été réussie.")
doc.add_page_break()

# III. Constats
doc.add_heading('III. Constats et réflexions personnels', level=1)
constats = """C’est avec un sentiment de satisfaction et de joie que j’ai eu à réaliser mon stage en responsabilité qui m’a permis non seulement de faire la présentation de plusieurs leçons mais aussi de mettre en œuvre les connaissances pédagogiques et didactiques acquises pendant la formation. Ce stage m’a permis de comprendre l’importance de préparer sa fiche de leçon de façon rigoureuse, de mettre en œuvre son cours en adoptant la méthode active et en utilisant des supports adéquats pour les activités prévues dans la présentation de la leçon. Globalement, les conditions dans lesquelles ce stage s’est déroulé me paraissent satisfaisantes. L’administration du CES/Gaweye1 et ma tutrice de stage, m’ont accompagné et aidé à mener mon stage en responsabilité dans des meilleures conditions."""
doc.add_paragraph(constats)
doc.add_page_break()

# Conclusion
doc.add_heading('Conclusion générale et recommandations', level=1)
conclusion = """Le stage en responsabilité permet non seulement au stagiaire de s’imprégner et de se familiariser avec l’atmosphère de la classe mais aussi de présenter des leçons tout en mettant en œuvre les connaissances pédagogiques et didactiques apprises tout au long de la formation. Mon stage en responsabilité s’est bien déroulé avec les soutiens de ma tutrice, de l’administration du CES/Gaweye1 et surtout de l’inspecteur Adamou Moussa et docteur Boubacar Maizoumbou qui m’ont apporté leurs expériences et expertises pour la réussite de mon stage en responsabilité. Les analyses et les conseils pratiques des superviseurs qui ont suivi mes différentes présentations des leçons ont été un cadre idéal d’échanges qui m’ont permis de relever mes forces et faiblesses à travers des observations pertinentes et objectives. Ces dernières m’aideront sans doute à bien mener ma carrière enseignante. Pour mieux améliorer le stage en responsabilité, je fais des recommandations suivantes :"""
doc.add_paragraph(conclusion)
doc.add_paragraph("À l’endroit de l’ENS :", style='List Bullet')
doc.add_paragraph("D’éviter de programmer le stage en Février qui est un mois de perturbation.", style='List Bullet')
doc.add_paragraph("De créer une école annexe pour ne plus amener les stagiaires dans les autres établissements afin d’éviter tout désagrément.", style='List Bullet')
doc.add_paragraph("D’accompagner les stagiaires le premier jour dans les établissements d’accueil pour qu’ils puissent avoir plus de considération.", style='List Bullet')
doc.add_paragraph("À l’endroit des stagiaires :", style='List Bullet')
doc.add_paragraph("De se soumettre aux exigences du stage en responsabilité.", style='List Bullet')
doc.add_paragraph("D’avoir plus de considération à l’égard des tuteurs de stage, de l’administration, des enseignants et des élèves de l’établissement d’accueil.", style='List Bullet')
doc.add_paragraph("À l’endroit des établissements d’accueil :", style='List Bullet')
doc.add_paragraph("D’afficher une grande considération à l’égard des stagiaires et de leur accorder une franche collaboration.", style='List Bullet')
doc.add_page_break()

# Références
doc.add_heading('Références bibliographiques', level=1)
refs = [
    "Administration du CES/Gaweye1 : tableau d’affichage dans le bureau du proviseur.",
    "Rapport de stage de ma tutrice, Ibrahim Hassane Nana Zoulaha, titre : Rapport de stage pour l’obtention de Master Professionnel au professorat de l’Enseignement Secondaire (MPPES).",
    "Internet, Le plan d’un rapport de stage, site : https://www.scribbr.fr"
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')
doc.add_page_break()

# Annexes
doc.add_heading('Annexes', level=1)
doc.add_paragraph("(Les fiches pédagogiques N°1 à N°6 sont disponibles dans le document original.)")
for i in range(1, 7):
    doc.add_paragraph(f"Fiche pédagogique N°0{i} – voir annexes du rapport original")

doc.save('Rapport_de_stage_MPPES2.docx')
print("✅ Fichier Word généré : Rapport_de_stage_MPPES2.docx")